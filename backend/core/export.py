"""Render a cover letter into a polished, downloadable document.

Two formats, one shared layout: a business-letter template with a sender
header (name + contact line), the date, an optional recipient block
(role / company), and the body split into paragraphs.

- `.docx` via python-docx
- `.pdf`  via reportlab (platypus)

Everything runs locally — no data leaves the device.
"""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from xml.sax.saxutils import escape


def _sender_name(sender: dict[str, str | None]) -> str:
    parts = [sender.get("name"), sender.get("surname")]
    return " ".join(p.strip() for p in parts if p and p.strip())


def _contact_line(sender: dict[str, str | None]) -> str:
    fields = [sender.get("email"), sender.get("phone"), sender.get("linkedin"), sender.get("github")]
    return "  ·  ".join(f.strip() for f in fields if f and f.strip())


def _today() -> str:
    return datetime.now(timezone.utc).strftime("%B %d, %Y")


def _paragraphs(letter: str) -> list[str]:
    """Split the letter into paragraphs on blank lines; keep single newlines."""
    blocks = [b.strip() for b in letter.replace("\r\n", "\n").split("\n\n")]
    return [b for b in blocks if b]


# ── DOCX ─────────────────────────────────────────────────────────────

def build_docx(letter: str, *, company: str | None, role: str | None, sender: dict[str, str | None]) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    name = _sender_name(sender)
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(16)

    contact = _contact_line(sender)
    if contact:
        p = doc.add_paragraph()
        run = p.add_run(contact)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph(_today())

    if company or role:
        doc.add_paragraph()
        recipient = doc.add_paragraph()
        if role:
            recipient.add_run(role)
            if company:
                recipient.add_run("\n")
        if company:
            recipient.add_run(company)

    doc.add_paragraph()
    for para in _paragraphs(letter):
        doc.add_paragraph(para)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF ──────────────────────────────────────────────────────────────

def build_pdf(letter: str, *, company: str | None, role: str | None, sender: dict[str, str | None]) -> bytes:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    base = getSampleStyleSheet()["Normal"]
    name_style = ParagraphStyle("Name", parent=base, fontName="Helvetica-Bold",
                                fontSize=18, leading=22, alignment=TA_LEFT, textColor=HexColor("#1A1330"))
    meta_style = ParagraphStyle("Meta", parent=base, fontSize=9, leading=13, textColor=HexColor("#666666"))
    recipient_style = ParagraphStyle("Recipient", parent=base, fontSize=11, leading=15, textColor=HexColor("#1A1330"))
    body_style = ParagraphStyle("Body", parent=base, fontSize=11, leading=16, spaceAfter=10, textColor=HexColor("#1A1330"))

    flow = []
    name = _sender_name(sender)
    if name:
        flow.append(Paragraph(escape(name), name_style))
    contact = _contact_line(sender)
    if contact:
        flow.append(Paragraph(escape(contact), meta_style))

    flow.append(Spacer(1, 16))
    flow.append(Paragraph(_today(), meta_style))

    if company or role:
        flow.append(Spacer(1, 12))
        if role:
            flow.append(Paragraph(escape(role), recipient_style))
        if company:
            flow.append(Paragraph(escape(company), recipient_style))

    flow.append(Spacer(1, 14))
    for para in _paragraphs(letter):
        flow.append(Paragraph(escape(para).replace("\n", "<br/>"), body_style))

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        topMargin=inch, bottomMargin=inch, leftMargin=inch, rightMargin=inch,
        title=f"Cover letter — {company}" if company else "Cover letter",
        author=name or "Cover Letter Local",
    )
    doc.build(flow)
    return buf.getvalue()


# ── dispatch ─────────────────────────────────────────────────────────

_MEDIA = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


def render(fmt: str, letter: str, *, company: str | None, role: str | None,
           sender: dict[str, str | None]) -> tuple[bytes, str]:
    """Return (file bytes, media type) for the requested format."""
    if fmt == "docx":
        return build_docx(letter, company=company, role=role, sender=sender), _MEDIA["docx"]
    if fmt == "pdf":
        return build_pdf(letter, company=company, role=role, sender=sender), _MEDIA["pdf"]
    raise ValueError(f"Unsupported export format: {fmt!r}")
